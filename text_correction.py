# text correction : version0
# (derived by line break in word boundaries)


# version0 problems so far
#    생성하는 word의 버전 차이로, 생성된 word파일이 렉이 걸리는 것으로 추정
#    오류가 아닌 부분도 붙이는 문제
#    시각장애인분들이 사용하려면, ui가 텍스트 선택가능해야하는데, 여기에서는 안됨
#    다양한 text incoding 혼재 시, 깨지는 텍스트들 생김
#    이미 잘 바뀌어 있는 pdf도 있음. 드래그 해서 긁으면 줄 단위로 끊기지만, pdf내에 text data는 제대로 잘 됨
#    text data가 pdf내에서 순서가 뒤죽박죽인 놈.



# version1 development may start after 8/6/2023 .
#    version1's goal is to migrate the code into compiler code languages, and also
#    enhance performance and add other features by ocr
#    모든 띄어쓰기 오류 고쳐주는 알고리즘 만들기
#    영어는 영어 고치는 툴 가져오기
#    solve all version0's problem
#    use deep learning & reinforced learning
#    good performance on non-textdata pdf too
#    data base for hangeul 만들기 -> 문법대로 정리해서
#    해당 os에 맞게 exe파일 만들어주는 shell파일 만들기
#    keep away the data from main file
#    make web version
#    docx to docx 도 만들기

# extra information
#    Why I used pdfminer.six, not PyPDF2: 
#        pdfminer.six automatically detect the text incoding of text in pdf 
#        In contrary, PyPDF2 can't handle multiple text incoding in pdf.
#    python-docx creates xml file that is compatible to microsoft word, not an exact microsoft word file

# 자카드 유사도(Jaccard Similarity)
# 자카드 유사도는 집합 간의 유사도를 계산하는 방법으로, 단어의 등장 여부에 기반합니다.
# 문장을 단어의 집합으로 변환한 후, 두 집합 간의 자카드 유사도를 계산합니다.
# 유사도 점수는 0부터 1까지의 값을 가지며, 1에 가까울수록 두 문장은 유사합니다.

# reference: https://github.com/dothinking/pdf2docx   ->   creating docx file, not readable file

# 사용자가 편한게 쓸수 있게 만들고
# 문법 적용해서 붙는 것들 띄기

# 오늘 할 것: 
# 목   차 > 목차
# 조사 데이터 제대로 정리하기
# 오류메시지 다듬기
# 숫자 없애는 기능
# 한다. 된다. 있다. 자주 쓰이는 표현 찾기
# 보편적 heuristic 찾기
# 코드 정리
# 패턴들 통합하고, 정리하고, 수정 및 추가 용이하게 만들어놓기

# 나중에 발전시켜야 하는 것
# 데이터 쌓기 -> 애초에 사전 데이터 이용


# developer mode: on / off
# - 고칠 때마다 고친것 표시
# - 통계 출력


# built-in library
import re
import tkinter as tk
from tkinter import filedialog
# from tkinter import font
# from tkinter import ttk
# import time
import os
import sys
import html

# import chardet

# need to be downloaded
from pdfminer.high_level import extract_text
from docx import Document
from screeninfo import get_monitors
from PIL import ImageTk, Image
# pip install pdfminer.six
# pip install python-docx
# pip install screeninfo
# pip install Pillow

# to make exe file,
# pip install pyinstaller
# for macOS: pyinstaller --noconsole --add-data "logo.jpg:." text_correction.py
# for Windows: pyinstaller --noconsole --add-data "logo.jpg;." text_correction.py

# pattern0 주변 띄어쓰기 인식해야하지 않을까(그 단어 한쪽은 적어도 space이게)



def open_pdf():
    file_path = filedialog.askopenfilename(filetypes=[("pdf 파일", "*.pdf")])
    if file_path:
        
        paths = file_path.split(sep='/')
        pdf_name = paths[len(paths)-1]
        file_name = pdf_name[:len(pdf_name)-4]
        
        text = pdf2text_pdfminerSix(file_path)
        
        save_path = file_path[:len(file_path)-len(pdf_name)-1]

        # don't need space correction
        text = correction_korean_special(text)
        save_file_docx(text,save_path,file_name)



def pdf2text_pdfminerSix(file_path):
    text = extract_text(file_path)
    return text
    
    
def type_text():
    def submit_button_clicked():
        text = input_entry.get()
        window_text.destroy()
        # window.destroy()
        
        # executable_path = sys.argv[0] # when execute .py file..?
        # executable_dir = os.path.dirname(os.path.abspath(executable_path)) # when execute .exe file..?
          
        if getattr(sys, 'frozen', False):
            # executable_dir = os.path.dirname(sys.executable)
            executable_dir = os.path.dirname(sys._MEIPASS)

        # 실행 파일이 포함되어 있지 않은 경우 (스크립트 모드)
        else:
            executable_dir = os.path.dirname(os.path.abspath(__file__))          
          
          
        text = correction_korean_special(text)
        text = correction_korean_space(text)
        save_file_docx(text,executable_dir,"result")
        # correction(text, "result", executable_path)

    
    # 윈도우 생성
    window_text = tk.Tk()

    # 텍스트 상자 생성
    input_entry = tk.Entry(window_text)
    input_entry.pack()

    # 제출 버튼 생성
    submit_button = tk.Button(window_text, text="제출", command=submit_button_clicked)
    submit_button.pack()

    window_text.mainloop()










def make_xml_compatible(string):
    # Replace control characters and NULL bytes
    modified_string = re.sub(r"[\x00-\x08\x0B-\x1F\x7F-\x84\x86-\x9F]", "\n", string)

    # Replace special characters with XML entities
    modified_string = html.escape(modified_string)

    return modified_string


def correction_korean_special(text):
    
    # 교육 지원 인력 자료작성 요령(2022) 규칙 3번, 10번, 12번, 13번, 14번, 15번, 22번
    # 유니코드 문자를 voice 변환 프로그램에서 인식 못하나?
    special_characters = { # text 직접 입력하려고 ctr1+c -> ctrl+v 하면 일부 문자들이 알파벳으로 변환 됨.. -> 감지 실패
    '\sI *\s': 1,
    'Ⅰ': "I", 'Ⅱ': "II", 'Ⅲ': "III", 'Ⅳ': "IV", 'Ⅴ': "V", 'Ⅵ': "VI",
    'Ⅶ': "VII", 'Ⅷ': "VIII", 'Ⅸ': "IX", 'Ⅹ': "X", 'Ⅺ': "XI", 'Ⅻ': "XII", 
    'Ⅼ': "L", 'Ⅽ': "C", 'Ⅾ': "D", 'Ⅿ': "M", 'ↀ': "CD", 'ↁ': 5000, 'ↂ': 10000,
    'ⅰ': 'i', 'ⅱ': 'ii', 'ⅲ': 'iii', 'ⅳ': 'iv', 'ⅴ': 'v', 'ⅵ': 'vi', 'ⅶ': 'vii', 
    'ⅷ': 'viii', 'ⅸ': 'ix', 'ⅹ': 'x', 'ⅺ': 'xi', 'ⅻ': 'xii',
    'ⅼ': 'l', 'ⅽ': 'c', 'ⅾ': 'd', 'ⅿ': 'm',
    '[(] *[)]': '----',
    '※': ' * ',
    '→': ' -> ',
    '…': '...',
    '+': ' + ', '=': ' = ', '×': ' * ', '÷': ' // ',
    '㎡': ' 제곱미터', '±': '플러스 마이너스',
    '⇒': '->',
    '㎝': 'cm',
    }
    
    
    jeongak = {
        "．": ".",
        "，": ",",
        "！": "!",
        "？": "?",
        "	": "    ",
        "＂": "\"",
        "＇": "\'",
    }
    
    for special in special_characters.keys():
        text = text.replace(str(special),str(special_characters[special])+"(고침)")
        # 왠지 모르겠는데, 전각기호가 반각기호로 변하면서 ' '+반각기호로 변함
        
    for special in jeongak.keys():
        text = text.replace(str(special),str(jeongak[special])+"(고침)")
        
    tmp = ''
    for j in range(len(text)-3):
        bangak = list(jeongak.values())
        # print(j)
        try: 
            if (text[j] == ' ') and (text[j+1] in bangak):
                text1 = text[:j]
                text2 = text[j+1:]
                text = text1 + text2
        except: # text의 길이가 계속 줄기에, 길이를 넘는 범위에 접근하면 끝내야한다.
            break
        
        
    return text


def correction_korean_space(text):
    
    
    word_freq = dict()
    p1 = re.compile('\w\w+')
    word_list = p1.findall(text)
    
    # print(word_list)
                
    for word in word_list:
        try:
            word_freq[word] += 1
        except:
            word_freq[word] = 1
            
        for i in range(2,len(word)+1):

            # 우리나라 단어의 대부분은 뒤에 부가적인 단어가 붙는다.(ex: 창의 -> 창의적, 창의성, 창의력, 창의에,....)
            # 아직 못 막는 것: 이기적인 -> '이기', '이기적', '이기적인' -> '이 기세를'이 '이기세를', '이론이 기술적인'이 '이론이기술적인
            piece_from_start = word[:i]
            if len(piece_from_start) > 1:
                try: word_freq[piece_from_start] += 1
                except: word_freq[piece_from_start] = 1


    tmp = sorted(word_freq.items(), key = lambda x: x[1], reverse=True)

    
    def english(word):
        return word.isalnum() and all(ord(letter) < 128 for letter in word)
        # return False
        
    # if english, 제외
    freq_3 = []
    for i in range(len(tmp)):
        if tmp[i][1] >= 7 and (not english(tmp[i][0])):
            freq_3.append(tmp[i][0])

        # 만약 길면 기준횟수 늘리기?
    

    patterns0 = freq_3 # 빈도 3회 이상의 단어들, (" 단 어 " -> ' 단어 ')
    
    # 각각의 패턴별로 예외적 단어 추가 가능하게? ex: ~에도 -> ' 에도 막부 ', '에도 시대' 이런 것 예외처리
    # 차라리 ' 한글자 '는 그런 유형 나오면 따로 처리? 앞으로도 붙여보고, 뒤로도 붙여보고 해서!
    
    # ' 단어 ' -> '단어 '
    patterns1 = [] 
    patterns1 += [' 은 ', ' 가 ', ' 께서 ', ' 에서 ', ' 을 ', ' 를 ', ' 에 ', ' 에게 ', ' 께 ', ' 한테 ', ' 더러 ']
    patterns1 += [' 에게서 ', ' 한테서 ', ' 로 ', ' 으로 ', ' 으로서 ', ' 의 ', ' 과 ', ' 와 ', ' 하고 ', ' 보다 ']
    patterns1 += [' 처럼 ', ' 만큼 ', ' 같이 ', ' 아 ', ' 야 ', ' 이여 ', ' 여 ', ' 이시여 ', ' 시여 ', ' 만 ', ]
    patterns1 += [' 조차 ', ' 마저 ', ' 까지부터 ', ' 이나 ', ' 나 ', ' 이나마 ', ' 나마 ', ' 이라도 ', ' 라도 ']
    patterns1 += [' 이야 ', ' 야 ', ' 이라야 ', ' 이 ', ' 인 ', ' 도 ', ' 는 ', 
                  ' 히 ', ' 적 ', ' 력 ', ' 에는 ']

    # ' 단어. ' -> '단어. '
    patterns1_1 = [' 다[.] ', ' 이다[.] ', ' 했다[.] ', ' 한다[.] ',' 하다[.] ', ' 였다[.] ', ' 랬다[.] ']
    
    # pattern1에 의해 붙은 '(앞단어)단어 ' -> '(앞단어) 단어 ' 해주는 예외처리 (앞단어에 의한 예외처리)
    patterns1_2 = [ '[.]이 ']

    patterns2 = [] # ' 단 어 ' -> ' 단어 '
    patterns2 += [' 그래서 ', ' 그러므로 ', ' 따라서 ', ' 때문에 ', ' 그러나 ', ' 그런데 ', ' 하지만 ', ' 인데도 ' ] 
    patterns2 += [ ' 한데도 ',  ' 불구하고 ', ' 오히려 ', ' 그리고 ', ' 또한 ', ' 게다가 ' , ' 또는 ', ' 아니면 ' ]
    patterns2 += [ ' 그나저나 ', ' 그래도 ' ]

    patterns3 = [] # '(word)단 어 ' --> '(word)단어 '
    patterns3 += ['때문에 ', '인데도 ','한데도 ','하지만 ', '으로 ']


    patterns4 = [] # ' 단 어(word)' --> ' 단어(word)'

    

        
    for pattern in patterns0:
        for i in range(1,len(pattern)):
            find = pattern[:i] + '\s' + pattern[i:]
            # replace = pattern[:i] + '((붙임0))' + pattern[i:]
            replace = pattern 
            replace += "(고침)"
            p = re.compile(find)
            text = p.sub(replace, text)


    for pattern in patterns1:
        find = '\s' + pattern[1:]
        # replace = '((붙임1))' + pattern[1:]
        replace = pattern[1:] 
        replace += "(고침)"
        text = re.sub(find, replace, text)
        
    for pattern in patterns1_1:
        find = '\s' + pattern[1:]
        # replace = '((붙임1))' + pattern[1:]
        replace = pattern[1:len(pattern)-4] + ". " 
        replace += "(고침)"
        text = re.sub(find, replace, text)
        
    for pattern in patterns1_2:
        find = pattern
        replace = ". " + pattern[3:]
        replace += "(고침-)"
        text = re.sub(find, replace, text)

    for pattern in patterns2:
        for i in range(2, len(pattern)-1):
            find = pattern[:i] + '\s' + pattern[i:]    
            # replace = pattern[:i] + '((붙임2))' + pattern[i:]  
            replace = pattern 
            replace += "(고침)"
            text = re.sub(find, replace, text)

    for pattern in patterns3:
        for i in range(1,len(pattern)-1):
            find = '(?P<front>\w+)' + pattern[:i] + '\s' + pattern[i:]
            # replace = '(\g<front>' + pattern[:i] + '((붙임3))' + pattern[i:]
            replace = '\g<front>' + pattern 
            replace += "(고침)"
            p = re.compile(find)
            text = p.sub(replace, text)


    # for pattern in patterns4.... 
            

    return text
    
    
def save_file_docx(text, save_path, file_name):


    text = make_xml_compatible(text)
    sentences = text.split(sep='\n')

    sentences2 = ''
    print(len(sentences))
    
    sentence_index = 0
    for sentence in sentences:
 
        if len(sentence) == 0 or sentence == ' ':
            continue

        
        end_of_sentence = sentence[len(sentence)-2:] == '. ' or sentence[len(sentence)-1] == '.'
        title_page = (not end_of_sentence) and sentence_index < 10
        small_titles = (not end_of_sentence) and len(sentence) < 18
        
        if end_of_sentence or title_page or small_titles:
            sentence += '\n'
            
        # sentences2 = sentences2 + '((띄어쓰기로))' + sentence
        sentences2 = sentences2 + sentence
        sentence_index += 1
    
    
    # doc.add_paragraph(sentences3)
    

    # 새로운 Word 문서 생성
    doc = Document()    
    for sentence in sentences2.split(sep='\n'):
        doc.add_paragraph(sentence)
        print("woho!!")


    x = 0
    existing_files = os.listdir(save_path)
    while True:
        if x == 0: save_name = "{}.docx".format(file_name)
        else: save_name = "{}_{}.docx".format(file_name,x)
        
        save_as = save_path + "/" + save_name
        
        if save_name in existing_files:
            x += 1
            continue

        doc.save(save_as)
        break



    complete_window = tk.Tk()  # 로딩 페이지를 나타내는 새로운 윈도우 생성
    complete_window.title("변환 완료!!!")
    
    try: complete_window.after(3000, complete_window.destroy)
    except: pass
    complete_window.mainloop()





# main
if __name__ == "__main__":
        
    # 윈도우 생성
    

    # 모든 모니터 정보 가져오기
    monitors = get_monitors()

    # 첫 번째 모니터 정보 사용하기
    primary_monitor = monitors[0]

    # 모니터의 가로 크기와 세로 크기 가져오기
    monitor_width = primary_monitor.width
    monitor_height = primary_monitor.height


    
    



    
    root = tk.Tk()
    root.overrideredirect(True)
    
    root_width = monitor_width// 4
    root_height = monitor_height// 3

    # 루트 윈도우를 모니터의 가운데로 배치
    root_pos_x = int((monitor_width - root_width) / 2)
    root_pos_y = int((monitor_height - root_height) / 2)
    
    root.geometry(f"{root_width}x{root_height}+{root_pos_x}+{root_pos_y}")

    # Load the image
    current_directory = os.path.abspath(__file__)
    print(current_directory)

    try:
        file_path = os.path.abspath(__file__)
        dir_path = os.path.dirname(file_path)        
        image = Image.open(dir_path + "/logo.jpg")
        image = image.resize((root_width, root_height))
    
    # Create a PhotoImage object from the image
        photo = ImageTk.PhotoImage(image)

    # Create a Label widget and assign the PhotoImage to its "image" parameter
        image_label = tk.Label(root, image=photo)
        image_label.pack()
    except:
        print("failed to find 'logo.jpg'")
    
    # time.sleep(3)
    
    def initialize_main():
        main = tk.Toplevel(root)
        
        # 루트 윈도우의 가로 크기와 세로 크기 설정
        main_width = monitor_width// 7
        main_height = monitor_height// 5

        # 루트 윈도우를 모니터의 가운데로 배치
        main_pos_x = int((monitor_width - main_width) / 2)
        main_pos_y = int((monitor_height - main_height) / 2)
        
        main.geometry(f"{main_width}x{main_height}+{main_pos_x}+{main_pos_y}")

        # "불러오기" 버튼 생성
        # main_font = font.Font(size=10)
        # button's size modified by the font size -> need to keep the size of button irrelevent to font size
        button1 = tk.Button(main, text="pdf 불러오기", command=open_pdf, width=int(main_width//10), height=int(main_height//33))
        button1.pack()
    
        button2 = tk.Button(main, text='text 직접 입력하기', command=type_text, width=int(main_width//10), height=int(main_height//33))
        button2.pack()
        
        main.protocol("WM_DELETE_WINDOW", root.destroy)
        main.mainloop()
        
        
    
    root.after(5000, initialize_main)
    
    root.mainloop()



# PDF 파일에서 텍스트 추출
